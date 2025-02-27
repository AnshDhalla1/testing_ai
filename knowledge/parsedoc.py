import os
import subprocess
import pypandoc
from docx import Document
import json

def convert_doc_to_docx(doc_path):
    """
    Converts a .doc file to .docx format using pandoc or libreoffice.
    Returns the new file path if conversion is successful.
    """
    doc_dir = os.path.dirname(doc_path)  # Get directory of the file
    doc_name = os.path.basename(doc_path)  # Extract filename
    docx_name = doc_name.replace(".doc", ".docx")  # Expected .docx filename
    docx_path = os.path.join(doc_dir, docx_name)  # Full .docx path

    subprocess.run(["libreoffice", "--headless", "--convert-to", "docx", doc_path, "--outdir", doc_dir])
    # try:
    #     # Convert using pypandoc
    #     pypandoc.convert_file(doc_path, "docx", outputfile=docx_path)
    # # except Exception:
    #     # Fallback: Use LibreOffice for conversion (save in same folder)
    #     subprocess.run(["libreoffice", "--headless", "--convert-to", "docx", doc_path, "--outdir", doc_dir])

    if os.path.exists(docx_path):
        print(f"Conversion successful: {docx_path}")
        return docx_path
    else:
        raise FileNotFoundError(f"Failed to convert {doc_path} to .docx")

def extract_text_from_doc(doc_path):
    """
    Extracts text and tables from a .docx file.
    """
    if doc_path.endswith(".doc"):  # Convert if it's a .doc file
        doc_path = convert_doc_to_docx(doc_path)

    doc = Document(doc_path)
    extracted_data = {"text": [], "tables": []}

    # Extract text from paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            extracted_data["text"].append(text)

    # Extract tables
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        extracted_data["tables"].append(table_data)

    return extracted_data

# doc_path = "/Users/Apple/Documents/Givery/data/JP resume format 021.doc"
# structured_output = extract_text_from_doc(doc_path)

# # Print structured output
# print(json.dumps(structured_output, ensure_ascii=False, indent=4))
